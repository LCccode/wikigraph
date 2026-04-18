"""Multi-format document conversion to markdown.

Converts Word, Excel, PDF, and plain text files to content.md.
Legacy .doc / .ppt formats are auto-converted via LibreOffice (soffice) if
available — graceful fallback otherwise.
PPT and image/video files require LLM Vision (handled in skill.md, not here).
"""

import shutil
import subprocess
import tempfile
from pathlib import Path


def _has_soffice() -> bool:
    """True if LibreOffice's `soffice` headless converter is on PATH."""
    return shutil.which("soffice") is not None or shutil.which("libreoffice") is not None


def _libreoffice_convert(src: Path, target_ext: str, timeout: int = 90) -> Path:
    """Convert src to target_ext via LibreOffice headless mode.

    Returns the path to the converted file. Raises on failure (caller
    should catch and decide whether to skip the file or fail the pipeline).

    target_ext: "docx" (for .doc) or "pptx" (for .ppt) etc.
    """
    if not _has_soffice():
        raise FileNotFoundError(
            "LibreOffice not installed. Install hints:\n"
            "  Ubuntu/Debian: apt install -y libreoffice-core\n"
            "  macOS:         brew install libreoffice\n"
            "  or download from https://www.libreoffice.org/download/"
        )
    bin_name = "soffice" if shutil.which("soffice") else "libreoffice"

    tmp_dir = Path(tempfile.mkdtemp(prefix="lcwiki_soffice_"))
    cmd = [
        bin_name, "--headless", "--convert-to", target_ext,
        str(src), "--outdir", str(tmp_dir),
    ]
    result = subprocess.run(cmd, capture_output=True, timeout=timeout)
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed for {src.name}: "
            f"{result.stderr.decode('utf-8', errors='replace')[:200]}"
        )
    converted = tmp_dir / f"{src.stem}.{target_ext}"
    if not converted.exists():
        # LibreOffice sometimes slightly alters the stem; find any matching file
        candidates = list(tmp_dir.glob(f"*.{target_ext}"))
        if candidates:
            converted = candidates[0]
        else:
            raise RuntimeError(f"LibreOffice produced no .{target_ext} output for {src.name}")
    return converted


def extract_pdf_text(path: Path) -> str:
    """Extract plain text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except Exception:
        return ""


def docx_to_markdown(path: Path, assets_dir: Path | None = None) -> tuple[str, list[Path]]:
    """Convert .docx to markdown using python-docx.

    Extracts paragraphs (with heading levels), tables, and images.
    Images are saved to assets_dir/images/ and referenced in markdown.

    Returns:
        (markdown_text, list_of_extracted_image_paths)
    """
    extracted_images: list[Path] = []
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        doc = Document(str(path))
        lines: list[str] = []

        # Build image rId → saved path mapping (extract all images first)
        img_map: dict[str, str] = {}  # rId → relative markdown path
        if assets_dir is not None:
            img_dir = assets_dir / "images"
            img_count = 0
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        img_data = rel.target_part.blob
                        content_type = rel.target_part.content_type
                        ext = ".png"
                        if "jpeg" in content_type or "jpg" in content_type:
                            ext = ".jpg"
                        elif "gif" in content_type:
                            ext = ".gif"
                        elif "bmp" in content_type:
                            ext = ".bmp"
                        img_count += 1
                        img_name = f"img-{img_count:03d}{ext}"
                        img_dir.mkdir(parents=True, exist_ok=True)
                        img_path = img_dir / img_name
                        img_path.write_bytes(img_data)
                        extracted_images.append(img_path)
                        img_map[rel.rId] = f"assets/images/{img_name}"
            except Exception:
                pass

        # Paragraphs — check each run for inline images
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()

            # Check if this paragraph contains inline images
            inline_imgs = []
            try:
                for run in para.runs:
                    drawing_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    if not drawing_elements:
                        drawing_elements = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
                    for drawing in drawing_elements:
                        # Find blip (image reference)
                        blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        for blip in blips:
                            embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed and embed in img_map:
                                inline_imgs.append(img_map[embed])
            except Exception:
                pass

            # Write paragraph text
            if not text and not inline_imgs:
                lines.append("")
                continue

            if text:
                if style.startswith("Heading 1"):
                    lines.append(f"# {text}")
                elif style.startswith("Heading 2"):
                    lines.append(f"## {text}")
                elif style.startswith("Heading 3"):
                    lines.append(f"### {text}")
                elif style.startswith("Heading 4"):
                    lines.append(f"#### {text}")
                elif style.startswith("List"):
                    lines.append(f"- {text}")
                else:
                    lines.append(text)

            # Insert image references right after the paragraph they belong to
            for img_ref in inline_imgs:
                img_name = Path(img_ref).stem
                lines.append(f"![{img_name}]({img_ref})")

        # Tables
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            lines.append("")
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            lines.extend([header, sep])
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
            lines.append("")

        # If there are images that weren't matched to paragraphs, list them at the end
        referenced = set(img_map.values())
        matched_in_text = set()
        for line in lines:
            for ref in referenced:
                if ref in line:
                    matched_in_text.add(ref)
        unmatched = [img for img in extracted_images if f"assets/images/{img.name}" not in matched_in_text]
        if unmatched:
            lines.append("")
            lines.append("## 其他文档图片")
            lines.append("")
            for img in unmatched:
                lines.append(f"![{img.stem}](assets/images/{img.name})")

        return "\n".join(lines), extracted_images
    except ImportError:
        return "", []
    except Exception:
        return "", []


def xlsx_to_markdown(path: Path) -> str:
    """Convert .xlsx to markdown using openpyxl.

    Each sheet becomes a ## section with a markdown table.
    """
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sections: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                if all(cell is None for cell in row):
                    continue
                rows.append([str(cell) if cell is not None else "" for cell in row])
            if not rows:
                continue

            sections.append(f"## Sheet: {sheet_name}")
            header = "| " + " | ".join(rows[0]) + " |"
            sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
            sections.extend([header, sep])
            for row in rows[1:]:
                sections.append("| " + " | ".join(row) + " |")
            sections.append("")

        wb.close()
        return "\n".join(sections)
    except ImportError:
        return ""
    except Exception:
        return ""


def convert_file(path: Path, assets_dir: Path | None = None) -> tuple[str, list[Path]]:
    """Convert a file to markdown content + extract assets (images).

    Args:
        path: Path to the source file.
        assets_dir: If provided, extract images to assets_dir/images/.
                    Images are referenced in markdown as ![](assets/images/img-001.png).
                    During compile, LLM can read these images with the Read tool.

    Returns:
        Tuple of (content_md: str, assets: list[Path]).

    Raises:
        ValueError if conversion fails or produces empty content.
    """
    ext = path.suffix.lower()
    assets: list[Path] = []

    # Legacy .doc / .ppt → auto-convert via LibreOffice, then recurse into the
    # modern .docx / .pptx branch. If LibreOffice is absent, raise a clear
    # ValueError (caller classifies as 'failed' and continues with other files).
    if ext == ".doc":
        if _has_soffice():
            try:
                converted = _libreoffice_convert(path, "docx")
                content, assets = docx_to_markdown(converted, assets_dir=assets_dir)
                content = f"<!-- converted from {path.name} via LibreOffice → docx -->\n\n{content}"
                if not content.strip():
                    raise ValueError(f"Conversion produced empty content for: {path.name}")
                return content, assets
            except Exception as e:
                raise ValueError(f".doc conversion failed: {e}")
        else:
            raise ValueError(
                f".doc requires LibreOffice (not installed). "
                f"Install: apt install -y libreoffice-core (Linux) or brew install libreoffice (macOS)"
            )
    if ext == ".ppt":
        if _has_soffice():
            try:
                converted = _libreoffice_convert(path, "pptx")
                # TODO: pptx_to_markdown not yet in lcwiki — for now, note the path
                content = (
                    f"<!-- converted from {path.name} via LibreOffice → pptx -->\n\n"
                    f"<!-- .pptx processing TBD, LibreOffice converted file at: {converted} -->\n"
                )
                return content, assets
            except Exception as e:
                raise ValueError(f".ppt conversion failed: {e}")
        else:
            raise ValueError(
                f".ppt requires LibreOffice (not installed). "
                f"Install: apt install -y libreoffice-core (Linux)"
            )

    if ext == ".docx":
        content, assets = docx_to_markdown(path, assets_dir=assets_dir)
    elif ext in (".xlsx", ".xls"):
        content = xlsx_to_markdown(path)
    elif ext == ".pdf":
        content = extract_pdf_text(path)
    elif ext in (".md", ".txt", ".rst"):
        content = path.read_text(encoding="utf-8", errors="replace")
    elif ext == ".pptx":
        content = f"<!-- PPT file: {path.name} - requires LLM Vision processing -->"
    elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"):
        content = f"<!-- Image file: {path.name} - requires LLM Vision processing -->"
    elif ext in (".mp4", ".mov", ".webm", ".mkv", ".avi", ".m4v", ".mp3", ".wav", ".m4a", ".ogg"):
        content = f"<!-- Audio/Video file: {path.name} - requires Whisper transcription -->"
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    if not content.strip():
        raise ValueError(f"Conversion produced empty content for: {path.name}")

    # Add source comment header
    content = f"<!-- converted from {path.name} -->\n\n{content}"

    return content, assets
